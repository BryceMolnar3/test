import re
import docx
from pymongo import MongoClient
import nltk
import gridfs

def extract_metadata(doc_path):
    """
    Extracts metadata from the first table of the document.
    For every row in the table, it uses the text in the first column as the attribute
    and the text in the second column as the value.
    
    Returns:
        A dictionary mapping each attribute to its corresponding value.
    """
    document = docx.Document(doc_path)
    if len(document.tables) < 1:
        raise ValueError("Document does not contain any tables.")
    
    table = document.tables[0]
    metadata = {}
    for row in table.rows:
        # Only process rows that have at least 2 cells.
        if len(row.cells) >= 2:
            key = row.cells[0].text.strip()
            value = row.cells[1].text.strip()
            metadata[key] = value
    return metadata

def extract_logical_middle_column_rows(doc_path, table_index=1):
    """
    Extracts the text from the middle column of the specified table (default: 2nd table)
    and groups consecutive physical rows into logical rows. For nonempty rows,
    consecutive rows with the same normalized text are merged; empty rows are kept as distinct.
    
    Returns:
        A list of logical rows (strings) from the middle column.
    """
    document = docx.Document(doc_path)
    
    if len(document.tables) <= table_index:
        raise ValueError(f"Document has only {len(document.tables)} table(s); cannot access table {table_index+1}.")
    
    table = document.tables[table_index]
    num_columns = len(table.columns)
    if num_columns == 0:
        raise ValueError("The selected table has no columns.")
    
    # For even number of columns, this picks the left-middle column.
    middle_index = num_columns // 2
    # Get all physical rows' text from the middle column.
    physical_rows = [row.cells[middle_index].text for row in table.rows]
    
    logical_rows = []
    last_normalized = None
    for cell_text in physical_rows:
        # Normalize by collapsing whitespace.
        normalized = " ".join(cell_text.split())
        # If the row is empty, always add it.
        if normalized == "":
            logical_rows.append(cell_text)
        else:
            # For nonempty rows, only add if different from the last nonempty row.
            if last_normalized is None or normalized != last_normalized:
                logical_rows.append(cell_text)
        last_normalized = normalized
    return logical_rows

def extract_logical_left_column_rows(doc_path, table_index=1):
    """
    Extracts the text from the middle column of the specified table (default: 2nd table)
    and groups consecutive physical rows into logical rows. For nonempty rows,
    consecutive rows with the same normalized text are merged; empty rows are kept as distinct.
    
    Returns:
        A list of logical rows (strings) from the middle column.
    """
    document = docx.Document(doc_path)
    
    if len(document.tables) <= table_index:
        raise ValueError(f"Document has only {len(document.tables)} table(s); cannot access table {table_index+1}.")
    
    table = document.tables[table_index]
    num_columns = len(table.columns)
    if num_columns == 0:
        raise ValueError("The selected table has no columns.")
    
    # For even number of columns, this picks the left-middle column.
    left_index = 0
    # Get all physical rows' text from the middle column.
    physical_rows = [row.cells[left_index].text for row in table.rows]
    physical_rows = physical_rows[6:-3]
    logical_rows = []
    
    for cell_text in physical_rows:
        # Normalize by collapsing whitespace.
        normalized = " ".join(cell_text.split())
        cell_num_and_text = cell_text.split(' ', 1)
        # If the row is empty, don't add it.
        if normalized != "":
            logical_rows.append((cell_num_and_text[0], cell_num_and_text[1]))

    return logical_rows


def split_text_into_verses(text):
    """
    Splits a block of text into verses by treating empty lines as delimiters.
    """
    verses = [verse.strip() for verse in re.split(r'\n\s*\n', text)]
    return verses

def match_verses(left_text, middle_text):
    """
    Matches each string in left_text to the closest string in middle_text using NLTK's edit distance.
    
    Args:
        middle_text (list of str): List of unnumbered text strings.
        left_text (list of tuples): List of (verse_number, text) tuples.
    
    Returns:
        list of tuples: [(assigned_verse_number, original_text), ...]
    """
    matched_text = []
    
    for text in middle_text:
        if text != "(omitted)":
            min_distance = float('inf')
            assigned_verse_number = None
            
            for verse_number, verse_text in left_text:
                distance = nltk.edit_distance(text, verse_text)
                
                if distance < min_distance:
                    min_distance = distance
                    assigned_verse_number = verse_number
            
            matched_text.append((assigned_verse_number, text))

    return matched_text

#Ian Local server: mongodb://127.0.0.1:27017
def send_to_mongodb(document_data, mongo_uri="mongodb://127.0.0.1:27017", db_name="document_db", collection_name="documents"):
    """
    Connects to MongoDB and inserts the document_data into the specified collection.
    """
    client = MongoClient(mongo_uri)
    db = client[db_name]
    fs = gridfs.GridFS(db)

    # Store the image and get its file ID
    with open(image_path, "rb") as f:
        image_id = fs.put(f, filename=image_path.split("/")[-1])
    
    # Add the image ID to your document
    document_data["image_file_id"] = image_id

    # Insert the updated document
    result = db[collection_name].insert_one(document_data)
    return result.inserted_id
    
    # collection = db[collection_name]
    # result = collection.insert_one(document_data)
    # return result.inserted_id

if __name__ == "__main__":
<<<<<<< HEAD:Data_Collection_Script/collectdata.py
    file_path = "01.docx"  # Change this to your filename if needed.
=======
    file_path = "test.docx"  # Change this to your filename if needed.
    image_path = "manuscript.png"
>>>>>>> collation:backend/collation/collectdata.py
    
    try:
        # Extract metadata from the first table.
        metadata = extract_metadata(file_path)
        
        # Extract logical rows from the middle column of the second table.
        middle_rows = extract_logical_middle_column_rows(file_path)
        left_rows = extract_logical_left_column_rows(file_path)
        
        # Determine which row contains the verses.
        if len(middle_rows) == 7:
            middle_verses = middle_rows[3]
        elif len(middle_rows) == 8:
            middle_verses = middle_rows[4]
        elif len(middle_rows) == 9:
            middle_verses = middle_rows[5]
        else:
            raise ValueError(f"Unexpected number of logical rows: {len(middle_rows)}. Expected 7, 8, or 9.")
        
        # Split the selected row's text into individual verses.
        verses = split_text_into_verses(middle_verses)
        matched_verses = match_verses(left_rows, verses)
        # Prepare the document to insert into MongoDB.
        document_data = {
            "filename": file_path,
            "metadata": metadata,
            "verses": matched_verses
        }
        
        # Send the document to MongoDB.
        inserted_id = send_to_mongodb(document_data)
        print(f"Data successfully inserted into MongoDB with ID: {inserted_id}\n")

        # Optionally, also print the extracted data.
        print("Extracted Manuscripts:")
        for i, verse in enumerate(left_rows, start=1):
            print(f"Verse {i}:\n{verse}\n")
        print("Extracted Verses:")
        for i, verse in enumerate(verses, start=1):
            print(f"Verse {i}:\n{verse}\n")
        print("Extracted Matches:")
        for i, verse in enumerate(matched_verses, start=1):
            print(f"Verse {i}:\n{verse}\n")
        
        print("Extracted Metadata:")
        for attr, value in metadata.items():
            print(f"{attr}: {value}")
        
    except Exception as e:
        print(f"Error: {e}")
